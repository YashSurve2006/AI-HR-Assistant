"""Resume analysis module — PDF/DOCX extraction, skill extraction, ATS scoring, and gap analysis."""

import os
import re
import zipfile
import xml.etree.ElementTree as ET
# pyrefly: ignore [missing-import]
import pdfplumber

# A comprehensive predefined dictionary of skills to extract
SKILLS_DB = [
    "Python", "Java", "JavaScript", "TypeScript", "React", "Node.js", "Express", "SQL", "MySQL",
    "PostgreSQL", "Oracle", "MongoDB", "HTML", "CSS", "Flask", "Django", "FastAPI", "Pandas", "NumPy",
    "Scikit-learn", "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "TensorFlow", "PyTorch",
    "Git", "GitHub", "Docker", "Kubernetes", "AWS", "Azure", "GCP", "CI/CD", "DevOps",
    "Linux", "Networking", "Cisco", "Excel", "Power BI", "Tableau",
    "C++", "C#", "Ruby", "PHP", "Swift", "Kotlin", "Go", "Rust", "GraphQL", "REST API",
    "Redux", "Next.js", "Vue.js", "Angular", "Tailwind CSS", "Bootstrap",
    "Agile", "Scrum", "Jira", "Leadership", "Communication", "Problem Solving"
]

# Standard high-demand industry skills benchmark for gap analysis
MARKET_DEMAND_SKILLS = [
    "Python", "SQL", "Git", "Docker", "AWS", "REST API", "JavaScript", "React", "Linux", "CI/CD", "Agile"
]

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file using python-docx or native zipfile/xml parsing."""
    try:
        # Try python-docx if installed
        try:
            import docx  # type: ignore
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.extend([cell.text for cell in row.cells if cell.text])
            if paragraphs:
                return "\n".join(paragraphs)
        except Exception:
            pass

        # Native DOCX parsing via zipfile & XML ElementTree
        text_pieces = []
        with zipfile.ZipFile(file_path) as z:
            xml_content = z.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            for elem in tree.iter():
                if elem.tag.endswith('t') and elem.text:
                    text_pieces.append(elem.text)
                elif elem.tag.endswith('p') or elem.tag.endswith('br'):
                    text_pieces.append('\n')
        return " ".join(text_pieces)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """Universal text extractor supporting PDF, DOCX, DOC, and TXT files."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""
    return ""

def preprocess_resume_text(text: str) -> str:
    """Clean and preprocess resume text."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_skills(text: str) -> list:
    """Extract technical and professional skills from text."""
    found_skills = set()
    
    for skill in SKILLS_DB:
        skill_lower = skill.lower()
        escaped_skill = re.escape(skill_lower)
        
        # Determine word boundaries depending on whether the skill starts/ends with alphanumeric
        prefix = r'\b' if re.match(r'^\w', skill_lower) else r'(?:\s|^)'
        suffix = r'\b' if re.match(r'\w$', skill_lower) else r'(?:\s|$|[.,;!?)])'
        
        pattern = f"{prefix}{escaped_skill}{suffix}"
        
        if re.search(pattern, text):
            found_skills.add(skill)
            
    return list(found_skills)

EXPERIENCE_KEYWORDS = [
    "experience", "work", "employment", "internship", "job", "career", "history",
    "years", "role", "position", "responsibilities", "duties", "tenure"
]
EDUCATION_KEYWORDS = [
    "education", "degree", "university", "college", "bachelor", "master", "phd",
    "btech", "mtech", "bs", "ms", "gpa", "academic", "institute", "diploma", "school"
]
PROJECTS_KEYWORDS = [
    "project", "developed", "built", "created", "led", "managed", "achieved",
    "award", "accomplishment", "certification", "certified", "implemented",
    "designed", "architected", "deployed", "optimized", "increased", "reduced"
]

def count_keywords(text: str, keywords: list) -> int:
    """Count occurrences of keywords in text using word boundaries where possible."""
    count = 0
    for kw in keywords:
        padded_text = f" {text} "
        count += padded_text.count(f" {kw} ")
    return count

def calculate_score(text: str, skills: list) -> dict:
    """Calculate resume score based on various components (0-100)."""
    # 1. Skills Score (max 25)
    skills_score = min(25, len(skills) * 3)
    
    # 2. Experience Score (max 25)
    exp_count = count_keywords(text, EXPERIENCE_KEYWORDS)
    experience_score = min(25, exp_count * 5)
    
    # 3. Education Score (max 20)
    edu_count = count_keywords(text, EDUCATION_KEYWORDS)
    education_score = min(20, edu_count * 5)
    
    # 4. Projects & Achievements Score (max 15)
    proj_count = count_keywords(text, PROJECTS_KEYWORDS)
    projects_score = min(15, proj_count * 3)
    
    # 5. Content Quality Score (max 15)
    length = len(text)
    if length > 2000:
        content_quality = 15
    elif length > 1000:
        content_quality = 12
    elif length > 500:
        content_quality = 8
    elif length > 200:
        content_quality = 5
    else:
        content_quality = 2
        
    final_score = skills_score + experience_score + education_score + projects_score + content_quality
    
    return {
        "score": final_score,
        "breakdown": {
            "skills": skills_score,
            "experience": experience_score,
            "education": education_score,
            "projects": projects_score,
            "content_quality": content_quality
        }
    }

def generate_insights(cleaned_text: str, skills: list, scoring: dict) -> dict:
    """Generate candidate strengths, weaknesses, missing in-demand skills, and actionable improvement suggestions."""
    strengths = []
    weaknesses = []
    suggestions = []
    
    bd = scoring.get("breakdown", {})
    score = scoring.get("score", 0)
    
    # Analyze Skills
    if len(skills) >= 6:
        strengths.append(f"Strong technical skill variety ({len(skills)} verified skills detected).")
    elif len(skills) >= 3:
        strengths.append(f"Solid foundational skillset ({', '.join(skills[:3])}).")
    else:
        weaknesses.append("Limited technical keywords detected, which may lower initial ATS filtering.")
        suggestions.append("Add a dedicated 'Technical Skills' section listing programming languages, tools, and frameworks.")

    # Missing in-demand market skills
    found_skills_lower = set(s.lower() for s in skills)
    missing_market_skills = [s for s in MARKET_DEMAND_SKILLS if s.lower() not in found_skills_lower]
    
    # Experience Analysis
    if bd.get("experience", 0) >= 15:
        strengths.append("Rich work experience articulation with industry-standard terminology.")
    else:
        weaknesses.append("Experience section could benefit from clearer role timelines and responsibilities.")
        suggestions.append("Clarify professional experience with clear job titles, company names, and bullet points.")

    # Projects & Quantifiable Achievements
    if bd.get("projects", 0) >= 10:
        strengths.append("Action-oriented project descriptions showcasing practical implementation.")
    else:
        weaknesses.append("Few quantifiable metrics or achievement action verbs found.")
        suggestions.append("Quantify your impact using metrics (e.g. 'Improved speed by 25%', 'Handled 10k+ users').")

    # Education & Credentials
    if bd.get("education", 0) >= 10:
        strengths.append("Clearly stated educational qualifications and academic background.")
    else:
        suggestions.append("Ensure your degree, institution name, and graduation year are explicitly listed.")

    # Overall Quality
    if bd.get("content_quality", 0) < 10:
        weaknesses.append("Resume content length is brief; some key details might be missing.")
        suggestions.append("Expand on project complexities, tools utilized, and team collaborations.")

    # Ensure at least 2 suggestions
    if len(suggestions) < 2:
        suggestions.append("Tailor your summary statement to highlight your primary domain expertise.")
        suggestions.append("Include relevant industry certifications (e.g. AWS, Scrum Master, Azure).")

    return {
        "strengths": strengths[:4],
        "weaknesses": weaknesses[:4] if weaknesses else ["Minor formatting and keyword optimization opportunities."],
        "missing_skills": missing_market_skills[:6],
        "suggestions": suggestions[:4]
    }

def analyze_resume(file_path: str) -> dict:
    """Analyze a resume (PDF/DOCX) and return extracted text details, skills, score, strengths, and suggestions."""
    text = extract_text_from_file(file_path)
    
    if not text.strip():
        return {
            "success": False,
            "error": "Could not extract text from the file. It may be image-based, encrypted, or empty."
        }
        
    cleaned_text = preprocess_resume_text(text)
    skills = extract_skills(cleaned_text)
    
    scoring = calculate_score(cleaned_text, skills)
    insights = generate_insights(cleaned_text, skills, scoring)
    
    return {
        "success": True,
        "filename": os.path.basename(file_path),
        "text_length": len(text),
        "skills": sorted(skills),
        "skill_count": len(skills),
        "score": scoring["score"],
        "score_breakdown": scoring["breakdown"],
        "strengths": insights["strengths"],
        "weaknesses": insights["weaknesses"],
        "missing_skills": insights["missing_skills"],
        "suggestions": insights["suggestions"]
    }
